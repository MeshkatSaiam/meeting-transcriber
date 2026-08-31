import os
import re
import sys
import json
import time
import argparse
import subprocess
import concurrent.futures
from datetime import datetime
from pathlib import Path
from dotenv import load_dotenv
import docx
from google import genai

# Ensure UTF-8 output for Bengali and Unicode characters in Windows terminal
if getattr(sys, "stdout", None) is not None and hasattr(sys.stdout, "encoding") and sys.stdout.encoding != "utf-8":
    try:
        sys.stdout.reconfigure(encoding="utf-8")
    except Exception:
        pass
if getattr(sys, "stderr", None) is not None and hasattr(sys.stderr, "encoding") and sys.stderr.encoding != "utf-8":
    try:
        sys.stderr.reconfigure(encoding="utf-8")
    except Exception:
        pass

DEFAULT_MODEL = "gemini-3.6-flash"
FALLBACK_MODEL = "gemini-3.5-flash-lite"
MAX_RETRIES_PER_MODEL = 3
MIN_CHUNK_DURATION = 3.0  # Minimum seconds required for a valid audio chunk

class TranscriptionCancelledException(Exception):
    """Raised when transcription is explicitly cancelled by the user."""
    pass

TRANSCRIPTION_PROMPT = """You are an expert verbatim meeting transcriber and diarization specialist.
Please provide an accurate, verbatim, diarized transcript of this audio recording.

Strict Transcription, Language Fidelity & Formatting Rules:

1. Verbatim Accuracy & No Invention:
   - Transcribe every word exactly as spoken. Do NOT paraphrase, summarize, omit, or invent words.
   - If a word or phrase is genuinely unintelligible or unclear due to background noise/overlap, mark it as "[inaudible]" instead of guessing.

2. Strict Language Fidelity (NO Translation):
   - Strictly preserve each utterance in its original spoken language — NEVER translate Bangla to English or English to Bangla, even for full sentences or paragraphs.
   - Spoken Bangla MUST be transcribed in native Bengali script (বাংলা). Never transcribe spoken Bengali into Romanized Banglish Latin letters.
   - Spoken English sentences MUST be transcribed in English script.
   - For code-mixed speech (Banglish), transcribe the Bengali words in Bengali script (বাংলা) while preserving embedded English technical terms, acronyms, brand names, and loanwords in English script (e.g. "আমরা এই PCB design টা Altium এ করবো").

3. Speaker Diarization & Labels:
   - Identify distinct speakers and label each speaker turn consistently (e.g., "[00:00] Speaker 1: ...", "[00:15] Speaker 2: ...").
   - Every change in speaker must start on a new line with its own timestamp and speaker label.

4. Mandatory Timestamp Format:
   - Use ONLY the single point format [MM:SS] (or [HH:MM:SS] if the recording exceeds one hour) at the start of each speaker turn.
   - Place exactly ONE timestamp at the start of each new speaker turn (e.g., "[00:00] Speaker 1: ...").
   - NEVER use timestamp ranges (e.g. do NOT use "[00:00 - 00:15]").
   - NEVER add spaces inside brackets (e.g. do NOT use "[ 00:00 ]").
   - NEVER use parentheses, unbracketed timestamps, or markdown bold around timestamps.

5. Output Format:
   - Output the transcript in chronological order.
"""

VOICE_REFERENCE_PROMPT = """You will be given REFERENCE audio clips, each labeled with a real person's name, followed by a MAIN recording to transcribe.

CRITICAL RULES — READ CAREFULLY:

1. Reference clips are NOT meeting content. They may sound similar or even identical to a portion of the main recording — this is expected and intentional, since they were deliberately extracted from voices you need to recognize. Do NOT transcribe reference clip audio into your output under any circumstance, even if it seems to repeat.

2. Your output transcript must contain ONLY content from the MAIN recording. Zero exceptions. If you find yourself about to write text describing what a reference clip says, stop — that content does not belong in your output.

3. Before writing your transcript, silently do this reasoning first (do not output this reasoning, just use it):
   a. Note the distinct voice characteristics (pitch, pace, tone) of each reference clip and its name.
   b. Listen through the entire main recording and identify each distinct speaker's voice.
   c. For each distinct speaker in the main recording, compare their voice against your notes from step (a).
   d. Only when a match is clear and confident, use that person's real name as the speaker label for ALL their turns throughout the recording — not just once, but consistently every time they speak.
   e. When labeling a matched speaker with their real name, use the EXACT name string provided above, with identical spelling — do not phonetically respell, transliterate differently, or alter the name in any way.
   f. If no clear match exists for a speaker, use generic labels (Speaker 1, Speaker 2, etc.) — never guess.

4. A common mistake to avoid: do not label only the first instance of a matched voice with their name and then drop back to a generic label for the rest of that same person's turns. If you identify a match, apply it consistently across the ENTIRE recording for that person.

Now transcribe the MAIN recording only, following the diarization and formatting rules below.
"""

VOICE_SAMPLES_DIR = Path("voice_samples")
VOICE_SAMPLES_INDEX = VOICE_SAMPLES_DIR / "samples.json"

def load_voice_samples() -> list[dict]:
    if not VOICE_SAMPLES_INDEX.exists():
        return []
    try:
        with open(VOICE_SAMPLES_INDEX, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return []

def transcribe_voice_sample_clip(clip_path: Path | str, model: str = DEFAULT_MODEL, api_key: str | None = None) -> str:
    """
    Transcribes a short voice sample clip once at add-time (plain verbatim transcription, no diarization).
    Cached in samples.json for deterministic sequence matching / sanity checks.
    """
    key = api_key
    if not key:
        load_dotenv()
        key = os.getenv("GEMINI_API_KEY")
    if not key:
        return ""

    src = Path(clip_path)
    if not src.exists():
        return ""

    try:
        client = genai.Client(api_key=key)
        uploaded = client.files.upload(file=str(src))
        while uploaded.state.name == "PROCESSING":
            time.sleep(0.5)
            uploaded = client.files.get(name=uploaded.name)

        prompt = "Transcribe this short audio clip verbatim in Bengali/English. Output ONLY the plain transcribed words, with no timestamps, no speaker labels, and no introductory or concluding text."
        response = client.models.generate_content(
            model=model,
            contents=[uploaded, prompt]
        )
        try:
            client.files.delete(name=uploaded.name)
        except Exception:
            pass

        return response.text.strip() if response.text else ""
    except Exception as e:
        print(f"[Voice Sample Cache Warning] Could not cache transcript for '{src.name}': {e}", file=sys.stderr)
        return ""

def get_sample_display_label(sample: dict) -> str:
    """
    Returns 'Name — Description' if description is present and non-empty, otherwise 'Name'.
    Enables duplicate names to be clearly distinguished at a glance.
    """
    name = (sample.get("name") or "Unknown Speaker").strip()
    desc = (sample.get("description") or "").strip()
    if desc:
        return f"{name} — {desc}"
    return name

def save_voice_sample(
    person_name: str,
    audio_file_path: Path | str,
    start_sec: float = 0.0,
    end_sec: float | None = None,
    description: str = ""
) -> dict:
    """
    Saves an audio clip into voice_samples/ capped to a maximum of 8.0 seconds via ffmpeg.
    Generates 16kHz mono MP3, transcribes once to cache transcript, stores metadata, and indexes in samples.json.
    """
    VOICE_SAMPLES_DIR.mkdir(parents=True, exist_ok=True)
    src_path = Path(audio_file_path)
    if not src_path.exists():
        raise FileNotFoundError(f"Source audio file '{src_path}' not found.")
    
    import uuid
    clean_name = sanitize_name_segment(person_name) or "Unknown_Person"
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    unique_suffix = uuid.uuid4().hex[:6]
    dest_filename = f"{clean_name}_{ts}_{unique_suffix}.mp3"
    dest_path = (VOICE_SAMPLES_DIR / dest_filename).resolve()

    # Enforce maximum 8.0 second length cap
    if end_sec is not None and end_sec > start_sec:
        clip_duration = min(float(end_sec - start_sec), 8.0)
    else:
        clip_duration = 8.0

    cmd = [
        "ffmpeg", "-y",
        "-ss", f"{max(0.0, float(start_sec)):.3f}",
        "-t", f"{clip_duration:.3f}",
        "-i", str(src_path),
        "-ar", "16000",
        "-ac", "1",
        str(dest_path)
    ]
    subprocess.run(cmd, check=True, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)

    actual_duration = get_audio_duration(dest_path)
    capped_duration = round(min(actual_duration, clip_duration), 2)
    dur_formatted = format_timestamp(capped_duration)

    # 1. Transcribe once at add-time to cache transcript in samples.json
    cached_transcript = transcribe_voice_sample_clip(dest_path)

    samples = load_voice_samples()
    entry = {
        "id": f"sample_{ts}_{unique_suffix}",
        "name": person_name.strip(),
        "description": description.strip(),
        "filename": dest_filename,
        "path": str(dest_path),
        "duration": dur_formatted,
        "duration_sec": capped_duration,
        "cached_transcript": cached_transcript,
        "include_in_transcription": True,
        "date_added": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    }
    samples.append(entry)
    with open(VOICE_SAMPLES_INDEX, "w", encoding="utf-8") as f:
        json.dump(samples, f, indent=2, ensure_ascii=False)
    return entry

def update_voice_sample_metadata(sample_id: str, new_name: str, new_description: str = "") -> bool:
    """
    Updates the name and description of a voice sample by its unique ID.
    """
    samples = load_voice_samples()
    found = False
    for s in samples:
        if s.get("id") == sample_id:
            s["name"] = new_name.strip()
            s["description"] = new_description.strip()
            found = True
            break
    if found:
        with open(VOICE_SAMPLES_INDEX, "w", encoding="utf-8") as f:
            json.dump(samples, f, indent=2, ensure_ascii=False)
    return found

def update_voice_sample_include(sample_id: str, include: bool) -> bool:
    samples = load_voice_samples()
    found = False
    for s in samples:
        if s.get("id") == sample_id:
            s["include_in_transcription"] = bool(include)
            found = True
            break
    if found:
        with open(VOICE_SAMPLES_INDEX, "w", encoding="utf-8") as f:
            json.dump(samples, f, indent=2, ensure_ascii=False)
    return found

def delete_voice_sample(sample_id: str) -> bool:
    samples = load_voice_samples()
    target = None
    remaining = []
    for s in samples:
        if s.get("id") == sample_id:
            target = s
        else:
            remaining.append(s)
    if target:
        try:
            fpath = Path(target.get("path", ""))
            if fpath.exists():
                fpath.unlink()
        except Exception:
            pass
        with open(VOICE_SAMPLES_INDEX, "w", encoding="utf-8") as f:
            json.dump(remaining, f, indent=2, ensure_ascii=False)
        return True
    return False

def extract_waveform_peaks(audio_path: Path | str, num_peaks: int = 350) -> list[float]:
    """
    Computes a simplified amplitude array (normalized 0.0 - 1.0) downsampled to num_peaks points.
    Uses pydub/numpy if available; falls back cleanly to ffmpeg raw PCM pipe.
    """
    src = Path(audio_path)
    if not src.exists():
        return [0.0] * num_peaks

    try:
        from pydub import AudioSegment
        import numpy as np
        seg = AudioSegment.from_file(str(src))
        seg = seg.set_channels(1).set_frame_rate(8000)
        samples = np.array(seg.get_array_of_samples(), dtype=np.float32)
        if len(samples) == 0:
            return [0.0] * num_peaks

        max_val = float(np.max(np.abs(samples)))
        if max_val > 0:
            samples = samples / max_val

        bucket_size = max(1, len(samples) // num_peaks)
        peaks = []
        for i in range(num_peaks):
            start = i * bucket_size
            end = min(len(samples), (i + 1) * bucket_size)
            if start < len(samples):
                chunk = samples[start:end]
                peak = float(np.max(np.abs(chunk))) if len(chunk) > 0 else 0.0
                peaks.append(round(peak, 3))
            else:
                peaks.append(0.0)
        return peaks

    except Exception:
        try:
            import numpy as np
            cmd = [
                "ffmpeg", "-y",
                "-i", str(src),
                "-ac", "1",
                "-ar", "8000",
                "-f", "s16le",
                "-"
            ]
            res = subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.DEVNULL, check=True)
            samples = np.frombuffer(res.stdout, dtype=np.int16).astype(np.float32)
            if len(samples) == 0:
                return [0.0] * num_peaks

            max_val = float(np.max(np.abs(samples)))
            if max_val > 0:
                samples = samples / max_val

            bucket_size = max(1, len(samples) // num_peaks)
            peaks = []
            for i in range(num_peaks):
                start = i * bucket_size
                end = min(len(samples), (i + 1) * bucket_size)
                if start < len(samples):
                    chunk = samples[start:end]
                    peak = float(np.max(np.abs(chunk))) if len(chunk) > 0 else 0.0
                    peaks.append(round(peak, 3))
                else:
                    peaks.append(0.0)
            return peaks
        except Exception as e:
            print(f"[Waveform Warning]: Could not extract peaks: {e}", file=sys.stderr)
            return [0.05] * num_peaks

RECONCILIATION_PROMPT = """You are an expert at multi-speaker dialogue analysis and meeting speaker diarization.
You are given transcripts from consecutive chronological chunks of a single recorded meeting.
In each chunk, speakers were labeled independently with local identifiers (e.g. "Speaker 1", "Speaker 2", "Speaker 3").

Your task is to reconcile speaker identities across all chunks to determine which local speaker in each chunk corresponds to the same physical person across the whole meeting.

Analysis guidelines:
1. Examine conversation continuity, subject matter, specific technical topics, and dialogue hand-offs across chunk boundaries.
2. Look for explicit speaker clues: self-introductions ("I am...", "Rahim here"), names used when addressing others ("Thanks Karim", "Sir"), job titles, and roles.
3. Compare speaking styles, linguistic habits, and turn-taking behavior.
4. Establish a unified global speaker labeling scheme: "Speaker 1", "Speaker 2", "Speaker 3", etc. (Optionally include identified names, e.g. "Speaker 1 (Rahim)" if clearly established).
5. If there is clear evidence linking a local speaker to a global speaker, set "confidence": "high".
6. If the identity is ambiguous, unconfirmed, or evidence is weak, set "confidence": "uncertain" and provide an explanatory note in "notes". DO NOT make unfounded guesses.
7. Output ONLY a valid JSON object matching this schema:
{
  "mappings": [
    {
      "chunk": 1,
      "local_speaker": "Speaker 1",
      "global_speaker": "Speaker 1",
      "confidence": "high",
      "notes": "Opened the meeting and introduced the team"
    }
  ]
}
"""

MEETING_NOTES_PROMPT = """You are an expert executive assistant and meeting intelligence analyst.
Analyze the ENTIRE merged transcript provided below (reading the full discussion from start to finish, not just the opening minutes or introduction, because company names, client identity, meeting purpose, and participant names often emerge partway through).

Extract structured naming metadata and generate comprehensive meeting notes:
1. company_name: The client, external company, organization, or institution discussed or represented (e.g. "Hameem Group", "Grameenphone", "Apex"). If no specific company or client is identifiable across the entire transcript, return null.
2. meeting_type: The type of meeting if it adds clarity (e.g. "Sales-Pitch", "Client-Review", "Technical-Discussion", "Interview", "Sprint-Planning", "Kickoff"). If generic or unclear, return null.
3. person_names: An array of key named individuals, participants, or clients mentioned (e.g. ["Rahim", "Karim"]). If none identifiable, return an empty array [].
4. topic_slug: A concise 2-3 word subject summary in English (letters/numbers/hyphens only, e.g. "Predictive-Maintenance" or "ERP-Migration").
5. notes_markdown: Comprehensive meeting minutes and notes formatted in clean Markdown.

Output ONLY a valid JSON object matching this exact schema:
{
  "company_name": "Company-Name-Or-Null",
  "meeting_type": "Meeting-Type-Or-Null",
  "person_names": ["Name1", "Name2"],
  "topic_slug": "2-3-Word-Topic-Slug",
  "notes_markdown": "## 1. Executive Summary\\n...\\n## 2. Key Decisions & Agreements\\n...\\n## 3. Action Items & Next Steps\\n...\\n## 4. Open Questions & Unresolved Topics\\n..."
}

Markdown Guidelines:
- ## 1. Executive Summary: concise summary of goals and outcomes across the entire meeting.
- ## 2. Key Decisions & Agreements: bulleted list of decisions, approvals, and policies.
- ## 3. Action Items & Next Steps: checklist with owners and deadlines (- [ ] **Task Description** — Owner: Name).
- ## 4. Open Questions & Unresolved Topics: list of blockers/follow-ups.
- Maintain bilingual accuracy: Bangla context in Bengali script (বাংলা) and English technical terms, tool names, acronyms, and product names as-is in English script.
- Be concise, professional, and strictly faithful to the actual transcript content.
"""

def format_timestamp(seconds: float) -> str:
    total_sec = int(round(seconds))
    hours = total_sec // 3600
    minutes = (total_sec % 3600) // 60
    secs = total_sec % 60
    if hours > 0:
        return f"{hours:02d}:{minutes:02d}:{secs:02d}"
    return f"{minutes:02d}:{secs:02d}"

def calculate_auto_chunk_minutes(total_duration_sec: float) -> float:
    """
    Automatic chunk sizing rule:
    - Under 20 min = single chunk
    - 20-60 min = 15-min chunks
    - 60-120 min = 25-min chunks
    - Over 120 min = 35-min chunks
    """
    total_min = total_duration_sec / 60.0
    if total_min < 20.0:
        return max(1.0, total_min + 1.0)
    elif total_min <= 60.0:
        return 15.0
    elif total_min <= 120.0:
        return 25.0
    else:
        return 35.0

def shift_timestamps_in_text(text: str, offset_seconds: float, chunk_index: int = None) -> str:
    chunk_label = f"Chunk {chunk_index}" if chunk_index is not None else "Transcript"
    offset_seconds = max(0.0, offset_seconds)

    def shift_time_str(t: str) -> str:
        parts = [int(p.strip()) for p in t.split(":")]
        if len(parts) == 2:
            total_sec = parts[0] * 60 + parts[1]
        elif len(parts) == 3:
            total_sec = parts[0] * 3600 + parts[1] * 60 + parts[2]
        else:
            return t
        return format_timestamp(total_sec + offset_seconds)

    # 1. Range pattern with arbitrary spacing or dashes: [ MM:SS - MM:SS ] or [MM:SS-MM:SS]
    range_pattern = r"\[\s*(\d{1,2}:\d{2}(?::\d{2})?)\s*[-–—to]+\s*(\d{1,2}:\d{2}(?::\d{2})?)\s*\]"
    text = re.sub(
        range_pattern,
        lambda m: f"[{shift_time_str(m.group(1))} - {shift_time_str(m.group(2))}]",
        text
    )

    # 2. Single timestamp pattern with optional internal spaces: [ MM:SS ] or [MM:SS]
    single_pattern = r"\[\s*(\d{1,2}:\d{2}(?::\d{2})?)\s*\]"
    text = re.sub(
        single_pattern,
        lambda m: f"[{shift_time_str(m.group(1))}]",
        text
    )

    # 3. Parentheses timestamps: ( MM:SS ) or (MM:SS)
    paren_pattern = r"\(\s*(\d{1,2}:\d{2}(?::\d{2})?)\s*\)"
    text = re.sub(
        paren_pattern,
        lambda m: f"[{shift_time_str(m.group(1))}]",
        text
    )

    # 4. Bolded timestamps: **00:15** or **[00:15]**
    bold_pattern = r"\*\*\s*\[?\s*(\d{1,2}:\d{2}(?::\d{2})?)\s*\]?\s*\*\*"
    text = re.sub(
        bold_pattern,
        lambda m: f"[{shift_time_str(m.group(1))}]",
        text
    )

    # Safety Net: Detect any unshifted, malformed, or unrecognized timestamp structures
    suspicious_bracket_pattern = r"\[([^\]\n]*\d+:\d+[^\]\n]*)\]"
    for match in re.finditer(suspicious_bracket_pattern, text):
        matched_str = match.group(0)
        if not re.match(r"^\[\d{2}:\d{2}(?::\d{2})?(?:\s*-\s*\d{2}:\d{2}(?::\d{2})?)?\]$", matched_str):
            print(
                f"\n[WARNING] {chunk_label}: Unrecognized timestamp format detected: '{matched_str}'. "
                f"Please verify this timestamp in the final document.",
                file=sys.stderr,
                flush=True
            )

    return text

def clean_chunk_text(text: str) -> str:
    lines = text.strip().splitlines()
    cleaned_lines = []
    for line in lines:
        stripped = line.strip()
        # Skip LLM intro / outro commentary
        if re.match(r"^(here is the|sure, here is|below is the|transcript:|---|\*\*\*|note:)", stripped, re.IGNORECASE):
            continue
        if stripped:
            cleaned_lines.append(stripped)
    return "\n\n".join(cleaned_lines)

def reconcile_speakers(client: genai.Client, chunks: list[dict], models: list[str], cancel_event=None, api_call_callback=None) -> dict[int, dict[str, dict]]:
    if cancel_event and cancel_event.is_set():
        raise TranscriptionCancelledException("Transcription cancelled by user.")
    valid_chunks = [c for c in chunks if not c.get("skipped") and c.get("transcript", "").strip()]
    if len(valid_chunks) <= 1:
        return {}

    print("\n" + "=" * 60)
    print("RECONCILING SPEAKER IDENTITIES ACROSS CHUNKS...")
    print("=" * 60, flush=True)

    prompt_sections = ["Here are the transcripts from each chronological chunk of the meeting:\n"]
    for c in valid_chunks:
        time_range = f"{format_timestamp(c['start_sec'])} - {format_timestamp(c['end_sec'])}"
        prompt_sections.append(f"--- CHUNK {c['chunk_index']} ({time_range}) ---\n{c['transcript']}\n")
    
    prompt_text = "\n".join(prompt_sections)

    for model in models:
        if cancel_event and cancel_event.is_set():
            raise TranscriptionCancelledException("Transcription cancelled by user.")
        try:
            if api_call_callback:
                try:
                    api_call_callback()
                except Exception:
                    pass
            response = client.models.generate_content(
                model=model,
                contents=[RECONCILIATION_PROMPT, prompt_text],
                config={"response_mime_type": "application/json"}
            )
            
            raw_json = response.text.strip()
            data = json.loads(raw_json)
            mappings_list = data.get("mappings", [])
            
            structured_mappings: dict[int, dict[str, dict]] = {}
            uncertain_count = 0

            print(f"\nSpeaker Reconciliation Table (using '{model}'):")
            for item in mappings_list:
                chunk_num = int(item.get("chunk", 1))
                local_spk = str(item.get("local_speaker", "")).strip()
                global_spk = str(item.get("global_speaker", "")).strip()
                confidence = str(item.get("confidence", "high")).strip().lower()
                notes = str(item.get("notes", "")).strip()

                if not local_spk or not global_spk:
                    continue

                if chunk_num not in structured_mappings:
                    structured_mappings[chunk_num] = {}

                structured_mappings[chunk_num][local_spk] = {
                    "global": global_spk,
                    "confidence": confidence,
                    "notes": notes
                }

                conf_str = "[UNCERTAIN - FLAGGED]" if confidence == "uncertain" else "(High confidence)"
                if confidence == "uncertain":
                    uncertain_count += 1
                print(f"  • Chunk {chunk_num}: '{local_spk}' -> '{global_spk}' {conf_str} | Reason: {notes}")

            if uncertain_count > 0:
                print(f"\n[WARNING] {uncertain_count} speaker mapping(s) were flagged as UNCERTAIN. Flagged in transcript output.", file=sys.stderr, flush=True)
            else:
                print("\nAll speaker mappings reconciled with high confidence.", flush=True)

            print("=" * 60 + "\n", flush=True)
            return structured_mappings

        except Exception as e:
            print(f"\n[Attempt failed for reconciliation with '{model}']: {e}", file=sys.stderr, flush=True)

    print("\n[WARNING] Speaker reconciliation could not be completed with available models. Proceeding with original labels.", file=sys.stderr, flush=True)
    return {}

def apply_speaker_mapping_to_chunk(chunk_text: str, chunk_mappings: dict[str, dict]) -> str:
    if not chunk_mappings:
        return chunk_text

    sorted_locals = sorted(chunk_mappings.keys(), key=lambda s: len(s), reverse=True)
    placeholder_map = {}
    temp_text = chunk_text

    for i, local_spk in enumerate(sorted_locals):
        placeholder = f"__SPK_MAP_PH_{i}__"
        info = chunk_mappings[local_spk]
        global_label = info["global"]
        
        if info.get("confidence", "").lower() == "uncertain":
            target_label = f"{global_label} (?)"
        else:
            target_label = global_label
            
        placeholder_map[placeholder] = target_label
        escaped_local = re.escape(local_spk)
        temp_text = re.sub(rf"\b{escaped_local}\b", placeholder, temp_text)

    for placeholder, target_label in placeholder_map.items():
        temp_text = temp_text.replace(placeholder, target_label)

    return temp_text

def get_audio_duration(audio_path: Path) -> float:
    cmd = [
        "ffprobe",
        "-v", "error",
        "-show_entries", "format=duration",
        "-of", "default=noprint_wrappers=1:nokey=1",
        str(audio_path)
    ]
    res = subprocess.run(cmd, capture_output=True, text=True, check=True)
    return float(res.stdout.strip())

def convert_to_mono_16k(input_path: Path, output_dir: Path) -> Path:
    output_dir.mkdir(parents=True, exist_ok=True)
    input_path = input_path.resolve()
    target_filename = f"{input_path.stem}.mp3"
    output_path = (output_dir / target_filename).resolve()

    # Avoid ffmpeg error when input and output resolve to the exact same file
    if input_path == output_path:
        target_filename = f"{input_path.stem}_converted.mp3"
        output_path = (output_dir / target_filename).resolve()

    print(f"Converting '{input_path.name}' to mono 16kHz MP3...")
    cmd = [
        "ffmpeg",
        "-y",
        "-i", str(input_path),
        "-ar", "16000",
        "-ac", "1",
        str(output_path)
    ]
    
    subprocess.run(cmd, check=True, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
    print(f"Converted master audio saved to: {output_path}")
    return output_path

def detect_silence_intervals(audio_path: Path, noise_db: str = "-30dB", min_duration: float = 0.4) -> list[tuple[float, float]]:
    print("Detecting silence intervals for smart chunk splitting...")
    cmd = [
        "ffmpeg",
        "-i", str(audio_path),
        "-af", f"silencedetect=noise={noise_db}:d={min_duration}",
        "-f", "null",
        "-"
    ]
    res = subprocess.run(cmd, capture_output=True, text=True)
    
    starts = [float(m.group(1)) for m in re.finditer(r"silence_start:\s*([0-9.]+)", res.stderr)]
    ends = [float(m.group(1)) for m in re.finditer(r"silence_end:\s*([0-9.]+)", res.stderr)]
    
    silences = []
    end_idx = 0
    for s in starts:
        while end_idx < len(ends) and ends[end_idx] < s:
            end_idx += 1
        if end_idx < len(ends):
            silences.append((s, ends[end_idx]))
            end_idx += 1
        else:
            silences.append((s, s + 0.5))
    return silences

def calculate_chunk_segments(total_duration: float, chunk_minutes: float, silences: list[tuple[float, float]]) -> list[tuple[float, float]]:
    chunk_target_sec = chunk_minutes * 60.0
    if total_duration <= chunk_target_sec:
        return [(0.0, total_duration)]

    silence_midpoints = [(s + e) / 2.0 for (s, e) in silences]
    segments = []
    cur_start = 0.0

    while cur_start + chunk_target_sec < total_duration:
        ideal_split = cur_start + chunk_target_sec
        if (total_duration - ideal_split) < MIN_CHUNK_DURATION:
            break

        window = min(chunk_target_sec * 0.25, 120.0)
        search_min = cur_start + max(chunk_target_sec * 0.5, chunk_target_sec - window)
        search_max = min(total_duration - MIN_CHUNK_DURATION, ideal_split + window)

        if search_min < search_max:
            candidates = [p for p in silence_midpoints if search_min <= p <= search_max]
        else:
            candidates = []

        if candidates:
            split_point = min(candidates, key=lambda p: abs(p - ideal_split))
            print(f"  - Chunk cut snapped to silence at {format_timestamp(split_point)} (ideal was {format_timestamp(ideal_split)})")
        else:
            split_point = ideal_split
            print(f"  - No silence found near window; cutting at ideal point {format_timestamp(split_point)}")

        if split_point - cur_start >= MIN_CHUNK_DURATION:
            segments.append((cur_start, split_point))
            cur_start = split_point
        else:
            break

    remainder = total_duration - cur_start
    if remainder >= MIN_CHUNK_DURATION:
        segments.append((cur_start, total_duration))
    elif segments:
        prev_start, _ = segments[-1]
        segments[-1] = (prev_start, total_duration)
    else:
        segments.append((0.0, total_duration))

    return segments

def split_audio_into_chunks(master_audio: Path, segments: list[tuple[float, float]], output_dir: Path) -> list[tuple[int, float, float, Path]]:
    chunks = []
    total_chunks = len(segments)

    if total_chunks == 1:
        return [(1, segments[0][0], segments[0][1], master_audio)]

    print(f"Splitting audio into {total_chunks} chunk(s)...")
    for i, (start_sec, end_sec) in enumerate(segments, 1):
        chunk_filename = f"{master_audio.stem}_chunk_{i:03d}.mp3"
        chunk_path = output_dir / chunk_filename
        
        cmd = [
            "ffmpeg",
            "-y",
            "-ss", f"{start_sec:.3f}",
            "-to", f"{end_sec:.3f}",
            "-i", str(master_audio),
            "-ar", "16000",
            "-ac", "1",
            str(chunk_path)
        ]
        subprocess.run(cmd, check=True, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
        chunks.append((i, start_sec, end_sec, chunk_path))
    
    return chunks

def create_combined_reference_chunk_audio(
    reference_samples: list[dict],
    chunk_audio_path: Path,
    output_path: Path
) -> tuple[float, list[dict]]:
    """
    Uses ffmpeg to concatenate reference samples (in a fixed, consistent sequence)
    followed by the chunk audio into ONE combined audio file.
    Tracks each sample's start/end offset and total_prefix_duration.
    """
    valid_refs = []
    current_offset = 0.0
    tracked_ranges = []

    for ref in reference_samples:
        p = Path(ref.get("path", ""))
        if p.exists():
            dur = ref.get("duration_sec") or get_audio_duration(p)
            start_t = round(current_offset, 2)
            end_t = round(current_offset + dur, 2)
            tracked_ranges.append({
                "name": ref.get("name", "Speaker"),
                "start_sec": start_t,
                "end_sec": end_t,
                "duration_sec": dur,
                "path": str(p),
                "cached_transcript": ref.get("cached_transcript", "")
            })
            valid_refs.append(p)
            current_offset = end_t

    total_prefix_duration = round(current_offset, 2)

    if not valid_refs:
        import shutil
        shutil.copy2(str(chunk_audio_path), str(output_path))
        return 0.0, []

    all_inputs = valid_refs + [chunk_audio_path]
    cmd = ["ffmpeg", "-y"]
    for inp in all_inputs:
        cmd.extend(["-i", str(inp)])

    filter_inputs = "".join(f"[{idx}:a]" for idx in range(len(all_inputs)))
    filter_str = f"{filter_inputs}concat=n={len(all_inputs)}:v=0:a=1[outa]"
    cmd.extend([
        "-filter_complex", filter_str,
        "-map", "[outa]",
        "-ar", "16000",
        "-ac", "1",
        str(output_path)
    ])

    subprocess.run(cmd, check=True, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
    return total_prefix_duration, tracked_ranges

def build_deterministic_reference_prompt(total_prefix_duration: float, tracked_ranges: list[dict]) -> str:
    """
    Builds boundary prompt stating exact reference sub-ranges, meeting start time,
    mandatory speaker naming dominance, generic fallback (never guess), verbatim fidelity,
    and no-translation rules.
    """
    range_descs = [
        f"seconds {r['start_sec']:.1f}–{r['end_sec']:.1f} is {r['name']}"
        for r in tracked_ranges
    ]
    ranges_str = ", ".join(range_descs)
    names_str = ", ".join(f"'{r['name']}'" for r in tracked_ranges)
    example_name = tracked_ranges[0]['name'] if tracked_ranges else "Person"

    prompt = (
        f"You will be given a combined audio recording where the first {total_prefix_duration:.1f} seconds are REFERENCE AUDIO CLIPS, "
        f"followed immediately by the MAIN MEETING RECORDING starting at {total_prefix_duration:.1f} seconds.\n\n"
        f"AUDIO STRUCTURE & BOUNDARIES (STRICTLY ENFORCED):\n"
        f"- Reference Portion: 0.0s to {total_prefix_duration:.1f}s contains voice reference clips ONLY ({ranges_str}).\n"
        f"- Main Meeting Portion: The actual meeting starts at {total_prefix_duration:.1f} seconds. Transcribe ONLY from {total_prefix_duration:.1f} seconds forward.\n"
        f"- Purpose: Use the reference portion solely to learn and recognize what these people's voices sound like ({names_str}).\n\n"
        f"CRITICAL SPEAKER DIARIZATION & TRANSCRIPTION RULES:\n\n"
        f"1. Zero Reference Transcription:\n"
        f"   - Do NOT transcribe any speech from the reference portion (0.0s to {total_prefix_duration:.1f}s) into your output under any circumstance.\n"
        f"   - Your output must contain ONLY speech from the main meeting portion.\n\n"
        f"2. Mandatory Speaker Naming (Dominant Rule):\n"
        f"   - When a speaking voice in the meeting matches one of the reference clips ({names_str}), YOU MUST label that speaker with their real name (e.g. '[00:00] {example_name}: ...') for ALL of their turns throughout the entire transcript.\n"
        f"   - When labeling a matched speaker with their real name, use the EXACT name string provided above, with identical spelling — do not phonetically respell, transliterate differently, or alter the name in any way.\n"
        f"   - Do NOT use generic labels like 'Speaker 1' for any voice that matches a reference clip.\n\n"
        f"3. Generic Fallback, Never Guess:\n"
        f"   - If and ONLY if a speaker in the meeting does NOT match any of the provided reference clips ({names_str}), label them with generic numbered speaker labels (e.g. 'Speaker 1', 'Speaker 2').\n"
        f"   - Never invent or guess real names for voices not included in the reference clips.\n\n"
        f"4. Verbatim Accuracy & No Invention:\n"
        f"   - Transcribe every spoken word verbatim. Do NOT paraphrase, summarize, omit, or invent words.\n"
        f"   - If speech is genuinely unclear due to overlap or noise, mark it as '[inaudible]' instead of guessing.\n\n"
        f"5. Strict Language Fidelity (NO Translation):\n"
        f"   - Strictly preserve each utterance in its original spoken language — NEVER translate Bangla to English or English to Bangla.\n"
        f"   - Spoken Bangla MUST be transcribed in native Bengali script (বাংলা), never in Romanized Banglish Latin letters.\n"
        f"   - Spoken English sentences MUST be transcribed in English script.\n"
        f"   - For code-mixed speech, transcribe Bengali words in Bengali script (বাংলা) while keeping embedded English technical terms, acronyms, brand names, and loanwords in English script.\n\n"
        f"6. Mandatory Timestamp Format:\n"
        f"   - Use ONLY the single point format [MM:SS] (or [HH:MM:SS] if over one hour) at the start of each speaker turn.\n"
        f"   - Place exactly ONE timestamp per speaker turn (e.g. '[00:00] {example_name}: ...').\n\n"
        f"7. Output the transcript in chronological order."
    )
    return prompt

def process_prefix_cutoff_and_shift(
    transcript: str,
    total_prefix_duration: float,
    chunk_duration: float = 0.0,
    tracked_ranges: list[dict] | None = None
) -> tuple[str, list[str]]:
    """
    Intelligently handles reference prefix cutoff and timestamp normalization:
    - If Gemini timestamped on the combined audio timeline (max timestamp > total_prefix_duration),
      drops lines < total_prefix_duration and shifts remaining timestamps by -total_prefix_duration.
    - If Gemini followed instructions to transcribe ONLY the meeting starting at [00:00],
      detects that content is meeting speech (via cached_transcript difference) and preserves the transcript.
    """
    if total_prefix_duration <= 0.0 or not transcript or not transcript.strip():
        return transcript, []

    timestamp_regex = re.compile(r"\[\s*(\d{1,2}):(\d{2})(?::(\d{2}))?\s*\]")

    def parse_time_sec(match) -> float:
        if match.group(3) is not None:
            return int(match.group(1)) * 3600 + int(match.group(2)) * 60 + int(match.group(3))
        return int(match.group(1)) * 60 + int(match.group(2))

    lines = transcript.splitlines()
    all_timestamps = []
    parsed_lines = []

    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        m = timestamp_regex.search(stripped)
        t_sec = parse_time_sec(m) if m else None
        if t_sec is not None:
            all_timestamps.append(t_sec)
        parsed_lines.append((stripped, t_sec, m))

    if not all_timestamps:
        return transcript, []

    max_ts = max(all_timestamps)
    # If max_ts > total_prefix_duration, the model clearly included the prefix timeline
    used_combined_timeline = (max_ts > total_prefix_duration + 0.5)

    remaining_lines = []
    dropped_lines = []

    if used_combined_timeline:
        # Case A: Combined Timeline -> Drop turns before total_prefix_duration, shift remaining
        in_prefix_region = False
        for stripped, t_sec, m in parsed_lines:
            if t_sec is not None:
                if t_sec < total_prefix_duration:
                    in_prefix_region = True
                    dropped_lines.append(stripped)
                else:
                    in_prefix_region = False
                    adjusted_sec = max(0.0, t_sec - total_prefix_duration)
                    new_ts = format_timestamp(adjusted_sec)
                    new_line = timestamp_regex.sub(f"[{new_ts}]", stripped, count=1)
                    remaining_lines.append(new_line)
            else:
                if in_prefix_region:
                    dropped_lines.append(stripped)
                else:
                    remaining_lines.append(stripped)
    else:
        # Case B: All timestamps are <= total_prefix_duration.
        # Check if the text matches the cached reference transcripts
        import difflib
        has_ref_match = False
        cached_texts = [r.get("cached_transcript", "").strip().lower() for r in (tracked_ranges or []) if r.get("cached_transcript")]

        for stripped, t_sec, m in parsed_lines:
            line_lower = stripped.lower()
            is_ref_line = False
            for cached in cached_texts:
                if len(cached) > 3:
                    sim = difflib.SequenceMatcher(None, line_lower, cached).ratio()
                    if sim >= 0.45 or cached in line_lower:
                        is_ref_line = True
                        has_ref_match = True
                        break
            if is_ref_line:
                dropped_lines.append(stripped)
            else:
                remaining_lines.append(stripped)

        # If no reference match was detected, ALL lines were meeting speech starting at [00:00]!
        if not has_ref_match:
            remaining_lines = [item[0] for item in parsed_lines]
            dropped_lines = []

    # Non-blocking log if reference lines are dropped
    if dropped_lines and tracked_ranges:
        import difflib
        dropped_text = " ".join(dropped_lines).lower()
        for ref in tracked_ranges:
            name = ref.get("name", "Unknown")
            cached = ref.get("cached_transcript", "").strip().lower()
            if cached and len(cached) > 3:
                sim = difflib.SequenceMatcher(None, dropped_text, cached).ratio()
                if sim < 0.35 and cached not in dropped_text and dropped_text not in cached:
                    print(
                        f"[Notice] Dropped {len(dropped_lines)} prefix line(s) for speaker reference.",
                        file=sys.stderr,
                        flush=True
                    )

    return "\n".join(remaining_lines).strip(), dropped_lines

def transcribe_chunk_with_fallback(
    client: genai.Client,
    uploaded_file,
    chunk_index: int,
    models: list[str],
    prompt: str = TRANSCRIPTION_PROMPT,
    max_retries: int = 3,
    cancel_event = None,
    api_call_callback = None
) -> tuple[str, str]:
    last_error = None
    contents = [uploaded_file, prompt]

    for model_index, model in enumerate(models):
        if cancel_event and cancel_event.is_set():
            raise TranscriptionCancelledException("Transcription cancelled by user.")

        if model_index > 0:
            print(f"[Chunk {chunk_index}] Primary model failed. Falling back to secondary model: '{model}'...", flush=True)

        for attempt in range(1, max_retries + 1):
            if cancel_event and cancel_event.is_set():
                raise TranscriptionCancelledException("Transcription cancelled by user.")
            try:
                if api_call_callback:
                    try:
                        api_call_callback()
                    except Exception:
                        pass
                response_stream = client.models.generate_content_stream(
                    model=model,
                    contents=contents
                )
                
                text_parts = []
                for chunk in response_stream:
                    if cancel_event and cancel_event.is_set():
                        raise TranscriptionCancelledException("Transcription cancelled by user.")
                    if chunk.text:
                        text_parts.append(chunk.text)
                
                full_text = "".join(text_parts).strip()
                return full_text, model

            except TranscriptionCancelledException:
                raise
            except Exception as e:
                last_error = e
                print(f"[Chunk {chunk_index}] Attempt {attempt}/{max_retries} failed for '{model}': {e}", file=sys.stderr, flush=True)

                err_msg = str(e).lower()
                if "generaterequestsperday" in err_msg or "perday" in err_msg or "404" in err_msg or "not_found" in err_msg or "no longer available" in err_msg:
                    print(f"[Chunk {chunk_index}] Model '{model}' unavailable or quota exhausted. Skipping retries for this model.", file=sys.stderr, flush=True)
                    break

                if attempt < max_retries:
                    default_backoff = max(15.0, 10.0 * attempt) if ("429" in err_msg or "resource_exhausted" in err_msg or "quota" in err_msg) else (2.0 ** attempt)
                    delay_match = re.search(r"retry in\s*([0-9.]+)\s*s", str(e), re.IGNORECASE) or re.search(r"retrydelay['\"]\s*:\s*['\"]?(\d+)", str(e), re.IGNORECASE)
                    if delay_match:
                        backoff_seconds = float(delay_match.group(1)) + 1.5
                    else:
                        backoff_seconds = default_backoff

                    print(f"[Chunk {chunk_index}] Rate limit / error encountered. Retrying in {backoff_seconds:.1f}s...", file=sys.stderr, flush=True)
                    # Check cancel periodically during sleep
                    sleep_start = time.time()
                    while time.time() - sleep_start < backoff_seconds:
                        if cancel_event and cancel_event.is_set():
                            raise TranscriptionCancelledException("Transcription cancelled by user.")
                        time.sleep(0.2)

    raise RuntimeError(f"Chunk {chunk_index} failed across all candidate models ({', '.join(models)}). Last error: {last_error}")

def process_chunk(chunk_info: tuple[int, float, float, Path], api_key: str, models: list[str], reference_samples: list[dict] | None = None, cancel_event = None, api_call_callback = None) -> dict:
    if cancel_event and cancel_event.is_set():
        raise TranscriptionCancelledException("Transcription cancelled by user.")
    chunk_index, start_sec, end_sec, chunk_file = chunk_info
    duration = end_sec - start_sec

    if duration < MIN_CHUNK_DURATION:
        print(f"[Chunk {chunk_index}] Skipping chunk {chunk_index} — negligible audio duration ({duration:.2f}s)", flush=True)
        return {
            "chunk_index": chunk_index,
            "start_sec": start_sec,
            "end_sec": end_sec,
            "transcript": "",
            "model_used": "SKIPPED",
            "chunk_file": chunk_file,
            "skipped": True
        }

    try:
        actual_duration = get_audio_duration(chunk_file)
        if actual_duration < MIN_CHUNK_DURATION or chunk_file.stat().st_size < 1024:
            print(f"[Chunk {chunk_index}] Skipping chunk {chunk_index} — negligible audio duration ({actual_duration:.2f}s)", flush=True)
            return {
                "chunk_index": chunk_index,
                "start_sec": start_sec,
                "end_sec": end_sec,
                "transcript": "",
                "model_used": "SKIPPED",
                "chunk_file": chunk_file,
                "skipped": True
            }
    except Exception:
        pass

    client = genai.Client(api_key=api_key)

    total_prefix_duration = 0.0
    tracked_ranges = []
    combined_chunk_file = None
    target_upload_file = chunk_file
    uploaded_file = None

    try:
        if cancel_event and cancel_event.is_set():
            raise TranscriptionCancelledException("Transcription cancelled by user.")
        # Step 2: If reference samples are provided, build single concatenated audio file
        active_refs = [r for r in (reference_samples or []) if Path(r.get("path", "")).exists()]
        if active_refs:
            combined_chunk_file = chunk_file.parent / f"combined_chunk_{chunk_index:03d}_{chunk_file.stem}.mp3"
            total_prefix_duration, tracked_ranges = create_combined_reference_chunk_audio(
                active_refs,
                chunk_file,
                combined_chunk_file
            )
            target_upload_file = combined_chunk_file
            prompt = build_deterministic_reference_prompt(total_prefix_duration, tracked_ranges)
            print(f"\n{'='*60}\n[CHUNK {chunk_index} PROMPT SENT TO GEMINI (VOICE SAMPLES: {len(active_refs)})]\n{'='*60}\n{prompt}\n{'='*60}\n", flush=True)
        else:
            prompt = TRANSCRIPTION_PROMPT
            print(f"\n{'='*60}\n[CHUNK {chunk_index} PROMPT SENT TO GEMINI (STANDARD)]\n{'='*60}\n{prompt}\n{'='*60}\n", flush=True)

        if cancel_event and cancel_event.is_set():
            raise TranscriptionCancelledException("Transcription cancelled by user.")

        print(f"[Chunk {chunk_index}] Uploading '{target_upload_file.name}'...", flush=True)
        uploaded_file = client.files.upload(file=str(target_upload_file))

        print(f"[Chunk {chunk_index}] Waiting for file to process...", flush=True)
        while uploaded_file.state.name == "PROCESSING":
            if cancel_event and cancel_event.is_set():
                raise TranscriptionCancelledException("Transcription cancelled by user.")
            time.sleep(0.5)
            uploaded_file = client.files.get(name=uploaded_file.name)

        if uploaded_file.state.name == "FAILED":
            raise RuntimeError(f"[Chunk {chunk_index}] Processing failed on File API: {uploaded_file.error}")

        print(f"[Chunk {chunk_index}] Transcribing...", flush=True)
        raw_transcript, model_used = transcribe_chunk_with_fallback(
            client=client,
            uploaded_file=uploaded_file,
            chunk_index=chunk_index,
            models=models,
            prompt=prompt,
            max_retries=MAX_RETRIES_PER_MODEL,
            cancel_event=cancel_event,
            api_call_callback=api_call_callback
        )

        # Step 4, 5, 6: Cut off prefix lines, shift timestamps, and run sanity checks
        if total_prefix_duration > 0.0:
            final_transcript, _ = process_prefix_cutoff_and_shift(
                transcript=raw_transcript,
                total_prefix_duration=total_prefix_duration,
                chunk_duration=duration,
                tracked_ranges=tracked_ranges
            )
        else:
            final_transcript = raw_transcript

        return {
            "chunk_index": chunk_index,
            "start_sec": start_sec,
            "end_sec": end_sec,
            "transcript": final_transcript,
            "model_used": model_used,
            "chunk_file": chunk_file,
            "skipped": False
        }
    finally:
        if uploaded_file:
            try:
                client.files.delete(name=uploaded_file.name)
            except Exception:
                pass
        if combined_chunk_file and combined_chunk_file.exists():
            try:
                combined_chunk_file.unlink()
            except Exception:
                pass

def sanitize_name_segment(seg: str) -> str:
    if not seg:
        return ""
    cleaned = re.sub(r"[^A-Za-z0-9]+", "-", str(seg).strip()).strip("-")
    parts = [p.capitalize() for p in cleaned.split("-") if p]
    return "-".join(parts)

def format_topic_slug(raw_slug: str) -> str:
    if not raw_slug:
        return ""
    cleaned = re.sub(r"[^A-Za-z0-9]+", "-", str(raw_slug).strip()).strip("-")
    parts = [p.capitalize() for p in cleaned.split("-") if p]
    if len(parts) > 3:
        parts = parts[:3]
    return "-".join(parts)

def build_meeting_base_name(
    date_str: str,
    company_name: str | None,
    meeting_type: str | None,
    person_names: list[str] | None,
    topic_slug: str | None,
    fallback_title: str = "Meeting"
) -> str:
    """
    Simplified naming scheme: Max 2 elements after the date.
    Priority 1: Company name
    Priority 2: Meeting type (e.g. Sales-Pitch, Client-Review, Technical-Discussion)
    Priority 3: Person name(s) (if no company)
    Priority 4: Topic slug (fallback)

    Rule: Date + highest priority element. Add a second element ONLY if needed for clarity.
    Never stack 3 or more elements.
    """
    clean_date = date_str or datetime.now().strftime("%Y-%m-%d")
    clean_company = sanitize_name_segment(company_name) if company_name and str(company_name).lower() != "null" else ""
    clean_type = sanitize_name_segment(meeting_type) if meeting_type and str(meeting_type).lower() != "null" else ""
    clean_slug = format_topic_slug(topic_slug) if topic_slug and str(topic_slug).lower() != "null" else ""

    clean_persons = []
    if person_names:
        for p in person_names:
            sp = sanitize_name_segment(p)
            if sp and sp not in clean_persons:
                clean_persons.append(sp)
    persons_str = "-".join(clean_persons[:2]) if clean_persons else ""

    # Priority 1: Company identified
    if clean_company:
        elem2 = clean_type if clean_type else clean_slug
        if elem2 and elem2.lower() not in clean_company.lower() and clean_company.lower() not in elem2.lower():
            elem2_short = "-".join(elem2.split("-")[:2])
            return f"{clean_date}_{clean_company}_{elem2_short}"
        return f"{clean_date}_{clean_company}"

    # Priority 2: Meeting type + Topic (no company)
    if clean_type and clean_slug:
        if clean_type.lower() in clean_slug.lower():
            return f"{clean_date}_{clean_slug}"
        slug_short = "-".join(clean_slug.split("-")[:2])
        return f"{clean_date}_{slug_short}_{clean_type}"

    # Priority 3: Person names (no company)
    if persons_str:
        elem2 = clean_type if clean_type else ("-".join(clean_slug.split("-")[:2]) if clean_slug else "")
        if elem2:
            return f"{clean_date}_{persons_str}_{elem2}"
        return f"{clean_date}_{persons_str}"

    # Priority 4: Topic slug
    if clean_slug:
        return f"{clean_date}_{clean_slug}"
    if clean_type:
        return f"{clean_date}_{clean_type}"

    # Fallback
    fallback = sanitize_name_segment(fallback_title) or "Meeting"
    return f"{clean_date}_{fallback}"

def get_file_recording_date(file_path: Path | str) -> str:
    """
    Determines recording date:
    1. Tries audio file metadata tags (ffprobe creation_time, date, etc.)
    2. Falls back to file creation/modification timestamp on disk
    Format: YYYY-MM-DD
    """
    p = Path(file_path)
    if not p.exists():
        return datetime.now().strftime("%Y-%m-%d")

    # 1. Try ffprobe metadata tags
    try:
        cmd = [
            "ffprobe",
            "-v", "error",
            "-show_entries", "format_tags",
            "-of", "json",
            str(p)
        ]
        res = subprocess.run(cmd, capture_output=True, text=True, check=True)
        data = json.loads(res.stdout)
        tags = data.get("format", {}).get("tags", {})
        for tag_key in ["creation_time", "date", "recorded_date", "ORIGINAL_RECORD_DATE", "com.apple.quicktime.creationdate"]:
            for k, v in tags.items():
                if k.lower() == tag_key.lower():
                    val = str(v).strip()
                    match = re.search(r"(\d{4})[-/.](\d{2})[-/.](\d{2})", val)
                    if match:
                        return f"{match.group(1)}-{match.group(2)}-{match.group(3)}"
    except Exception:
        pass

    # 2. Fall back to file disk timestamp
    try:
        stat = p.stat()
        time_sec = min(stat.st_mtime, getattr(stat, "st_ctime", stat.st_mtime))
        return datetime.fromtimestamp(time_sec).strftime("%Y-%m-%d")
    except Exception:
        return datetime.now().strftime("%Y-%m-%d")

def generate_meeting_notes(transcript_text: str, model: str = DEFAULT_MODEL, models_to_try: list[str] = None, api_key: str = None, cancel_event = None, api_call_callback = None) -> dict:
    """
    Generates structured meeting notes, company name, meeting type, person names, and topic slug from transcript text via Gemini.
    Returns: {
        "notes": str,
        "topic_slug": str,
        "company_name": str,
        "meeting_type": str,
        "person_names": list[str]
    }
    """
    if cancel_event and cancel_event.is_set():
        raise TranscriptionCancelledException("Transcription cancelled by user.")

    if not transcript_text or not transcript_text.strip():
        raise ValueError("Transcript text is empty. Cannot generate meeting notes.")

    if not api_key:
        load_dotenv()
        api_key = os.getenv("GEMINI_API_KEY")
    if not api_key:
        raise ValueError("GEMINI_API_KEY not found in .env file.")

    client = genai.Client(api_key=api_key)

    if models_to_try is None:
        models_to_try = [model]
        if FALLBACK_MODEL not in models_to_try:
            models_to_try.append(FALLBACK_MODEL)

    prompt = f"{MEETING_NOTES_PROMPT}\n\n--- FULL MERGED MEETING TRANSCRIPT ---\n{transcript_text}"

    last_error = None
    for m in models_to_try:
        if cancel_event and cancel_event.is_set():
            raise TranscriptionCancelledException("Transcription cancelled by user.")
        try:
            if api_call_callback:
                try:
                    api_call_callback()
                except Exception:
                    pass
            response = client.models.generate_content(
                model=m,
                contents=prompt,
                config={"response_mime_type": "application/json"}
            )
            raw_text = response.text.strip()
            try:
                data = json.loads(raw_text)
                notes = data.get("notes_markdown", "").strip()
                raw_slug = data.get("topic_slug", "").strip()
                formatted_slug = format_topic_slug(raw_slug)
                raw_company = data.get("company_name", "")
                company_str = sanitize_name_segment(raw_company) if raw_company and str(raw_company).lower() != "null" else ""
                raw_type = data.get("meeting_type", "")
                type_str = sanitize_name_segment(raw_type) if raw_type and str(raw_type).lower() != "null" else ""

                raw_persons = data.get("person_names", [])
                clean_persons = []
                if isinstance(raw_persons, list):
                    for p in raw_persons:
                        sp = sanitize_name_segment(p)
                        if sp and sp not in clean_persons:
                            clean_persons.append(sp)

                if notes:
                    return {
                        "notes": notes,
                        "topic_slug": formatted_slug,
                        "company_name": company_str,
                        "meeting_type": type_str,
                        "person_names": clean_persons
                    }
            except Exception:
                if raw_text:
                    return {
                        "notes": raw_text,
                        "topic_slug": "",
                        "company_name": "",
                        "meeting_type": "",
                        "person_names": []
                    }
        except Exception as e:
            last_error = e
            print(f"[Meeting notes attempt failed for '{m}']: {e}", file=sys.stderr, flush=True)

    raise RuntimeError(f"Failed to generate meeting notes across candidate models ({', '.join(models_to_try)}). Last error: {last_error}")

def extract_transcript_from_file(file_path: Path) -> str:
    """
    Extracts text from a previously saved .docx or .txt file.
    """
    file_path = Path(file_path)
    if not file_path.exists():
        raise FileNotFoundError(f"File '{file_path}' does not exist.")

    if file_path.suffix.lower() == ".docx":
        doc = docx.Document(str(file_path))
        lines = []
        for p in doc.paragraphs:
            txt = p.text.strip()
            if txt:
                lines.append(txt)
        return "\n\n".join(lines)
    else:
        try:
            return file_path.read_text(encoding="utf-8")
        except UnicodeDecodeError:
            return file_path.read_text(encoding="cp1252", errors="replace")

def upload_to_google_drive(file_path: Path, log_fn=None) -> bool:
    """
    Stub for Google Drive upload.
    Will be populated once OAuth credentials are configured.
    """
    def log(msg: str):
        if log_fn:
            log_fn(msg)
        print(msg, flush=True)

    log(f"[Google Drive] Auto-upload requested for '{file_path.name}'.")
    log("[Google Drive] NOTE: OAuth upload stub active. File is saved locally. (Drive credentials pending).")
    return True

def delete_intermediate_files(files: list[Path], log_fn=None) -> int:
    """
    Deletes intermediate converted master audio and chunk files upon successful pipeline completion.
    Never deletes the original input file.
    """
    deleted_count = 0
    for f in files:
        try:
            if f and f.exists() and f.is_file():
                f.unlink()
                deleted_count += 1
        except Exception as e:
            if log_fn:
                log_fn(f"[Warning] Could not remove temp file '{f.name}': {e}")

    if deleted_count > 0 and log_fn:
        log_fn(f"[Cleanup] Deleted {deleted_count} intermediate audio file(s) (converted master & chunks).")

    return deleted_count

def save_transcript_docx(
    output_path: Path,
    title: str,
    merged_transcript: str,
    meeting_notes: str | None = None,
    metadata: dict | None = None
) -> Path:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc = docx.Document()
    
    # Document Title
    has_notes = bool(meeting_notes and meeting_notes.strip())
    has_transcript = bool(merged_transcript and merged_transcript.strip())

    if has_notes and not has_transcript:
        doc.add_heading(f"Meeting Notes & Action Items: {title}", level=1)
    elif has_transcript and not has_notes:
        doc.add_heading(f"Meeting Diarized Transcript: {title}", level=1)
    else:
        doc.add_heading(f"Meeting Transcript & Notes: {title}", level=1)
    
    # Metadata info
    if metadata:
        meta_p = doc.add_paragraph()
        meta_items = []
        if "duration" in metadata:
            meta_items.append(f"Duration: {metadata['duration']}")
        if "model" in metadata:
            meta_items.append(f"Model: {metadata['model']}")
        if "date" in metadata:
            meta_items.append(f"Generated: {metadata['date']}")
        meta_p.add_run(" | ".join(meta_items)).italic = True

    # 1. Meeting Notes Section (if generated)
    if meeting_notes and meeting_notes.strip():
        doc.add_heading("Meeting Notes & Action Items", level=2)
        for block in meeting_notes.strip().split("\n\n"):
            block = block.strip()
            if not block:
                continue
            if block.startswith("## "):
                doc.add_heading(block.lstrip("#").strip(), level=3)
            elif block.startswith("# "):
                doc.add_heading(block.lstrip("#").strip(), level=2)
            else:
                doc.add_paragraph(block)

    # 2. Diarized Transcript Section
    if merged_transcript and merged_transcript.strip():
        doc.add_heading("Full Diarized Transcript", level=2)
        for block in merged_transcript.strip().split("\n\n"):
            block = block.strip()
            if not block:
                continue
            p = doc.add_paragraph()
            # Match speaker / timestamp prefix to bold
            match = re.match(r"^(\*{0,2}\[?\d{1,2}:\d{2}(?::\d{2})?\]?\*{0,2}\s*(?:Speaker\s+[^\n:]+|[A-Za-z\u0980-\u09FF\s\(\)\?]+)?:?\s*)(.*)$", block, re.IGNORECASE | re.DOTALL)
            if match and match.group(1).strip():
                prefix = match.group(1).replace("**", "")
                rest = match.group(2)
                run_prefix = p.add_run(prefix)
                run_prefix.bold = True
                p.add_run(rest)
            else:
                p.add_run(block)

    doc.save(str(output_path))
    return output_path

def run_transcription_pipeline(
    audio_path: str | Path,
    model: str = DEFAULT_MODEL,
    chunk_minutes: float | str | None = None,
    max_workers: int | None = None,
    auto_save: bool = False,
    reference_samples: list[dict] | None = None,
    log_callback = None,
    status_callback = None,
    cancel_event = None,
    api_call_callback = None
) -> dict:
    """
    Executes chunking, parallel transcription, speaker reconciliation, and timestamp globalization.
    Returns structured results dict without forcing immediate save if auto_save=False.
    """
    def log(msg: str):
        if log_callback:
            try:
                log_callback(msg)
            except Exception:
                pass
        print(msg, flush=True)

    def set_status(action: str, current: int = 0, total: int = 1):
        if status_callback:
            try:
                status_callback(action, current, total)
            except Exception:
                pass

    if cancel_event and cancel_event.is_set():
        raise TranscriptionCancelledException("Transcription cancelled by user.")

    input_file = Path(audio_path)
    if not input_file.exists():
        raise FileNotFoundError(f"Input file '{input_file}' does not exist.")

    load_dotenv()
    api_key = os.getenv("GEMINI_API_KEY")
    if not api_key:
        raise ValueError("GEMINI_API_KEY not found in .env file.")

    models_to_try = [model]
    if FALLBACK_MODEL not in models_to_try:
        models_to_try.append(FALLBACK_MODEL)

    converted_dir = Path("converted")
    output_dir = Path("output")

    set_status("Converting audio to 16kHz mono MP3...", 0, 1)
    master_audio = convert_to_mono_16k(input_file, converted_dir)

    if cancel_event and cancel_event.is_set():
        raise TranscriptionCancelledException("Transcription cancelled by user.")

    total_duration = get_audio_duration(master_audio)
    
    # Resolve chunk size (auto vs manual)
    if chunk_minutes is None or chunk_minutes == "auto" or (isinstance(chunk_minutes, (int, float)) and chunk_minutes <= 0):
        effective_chunk_mins = calculate_auto_chunk_minutes(total_duration)
        chunk_mode_str = f"Auto ({effective_chunk_mins:.1f} mins based on {format_timestamp(total_duration)} audio)"
    else:
        effective_chunk_mins = float(chunk_minutes)
        chunk_mode_str = f"Manual ({effective_chunk_mins:.1f} mins)"

    log(f"Total audio duration: {format_timestamp(total_duration)} ({total_duration:.2f}s)")
    log(f"Chunk sizing mode: {chunk_mode_str}")
    log(f"Primary model: '{model}' (Fallback: '{FALLBACK_MODEL}')")
    if reference_samples:
        sample_names = ", ".join([s.get("name", "Unknown") for s in reference_samples])
        log(f"Voice reference matching active with {len(reference_samples)} sample(s): [{sample_names}]")

    if cancel_event and cancel_event.is_set():
        raise TranscriptionCancelledException("Transcription cancelled by user.")

    set_status("Detecting silence intervals...", 0, 1)
    silences = detect_silence_intervals(master_audio)
    segments = calculate_chunk_segments(total_duration, effective_chunk_mins, silences)

    if cancel_event and cancel_event.is_set():
        raise TranscriptionCancelledException("Transcription cancelled by user.")

    set_status("Splitting audio into chunks...", 0, len(segments))
    chunk_items = split_audio_into_chunks(master_audio, segments, converted_dir)

    total_chunks = len(chunk_items)
    log(f"\nStarting concurrent transcription for {total_chunks} chunk(s)...\n")
    set_status(f"Starting transcription for {total_chunks} chunk(s)...", 0, total_chunks)

    completed_results = []
    workers = max_workers or min(total_chunks, 8)
    with concurrent.futures.ThreadPoolExecutor(max_workers=workers) as executor:
        future_map = {
            executor.submit(process_chunk, chunk_info, api_key, models_to_try, reference_samples, cancel_event, api_call_callback): chunk_info[0]
            for chunk_info in chunk_items
        }

        for future in concurrent.futures.as_completed(future_map):
            if cancel_event and cancel_event.is_set():
                raise TranscriptionCancelledException("Transcription cancelled by user.")
            chunk_idx = future_map[future]
            try:
                result = future.result()
                completed_results.append(result)
                time_range = f"{format_timestamp(result['start_sec'])} - {format_timestamp(result['end_sec'])}"
                
                clock_now = datetime.now().strftime("%H:%M:%S")
                done_count = len(completed_results)
                if result.get("skipped"):
                    log(f"[{clock_now}] Chunk {result['chunk_index']} skipped (negligible duration).")
                else:
                    log(f"[{clock_now}] Chunk {result['chunk_index']}/{total_chunks} completed ({time_range}) [Model: {result['model_used']}]")
                
                set_status(f"Transcribing chunks ({done_count} of {total_chunks} completed)...", done_count, total_chunks)
            except TranscriptionCancelledException:
                raise
            except Exception as exc:
                log(f"\n[ERROR] Chunk {chunk_idx} generated an exception: {exc}")
                raise exc

    if cancel_event and cancel_event.is_set():
        raise TranscriptionCancelledException("Transcription cancelled by user.")

    # Sort chunks chronologically by chunk index
    completed_results.sort(key=lambda r: r["chunk_index"])

    # Speaker Reconciliation across chunks
    set_status("Reconciling speaker identities across chunks...", total_chunks, total_chunks)
    client = genai.Client(api_key=api_key)
    speaker_mappings = reconcile_speakers(client, completed_results, models=models_to_try, cancel_event=cancel_event, api_call_callback=api_call_callback)

    # Globalize timestamps, apply reconciled speaker labels, and merge
    merged_sections = []
    models_used_set = set()

    for r in completed_results:
        if r.get("skipped") or not r["transcript"].strip():
            continue
        models_used_set.add(r["model_used"])
        cleaned_text = clean_chunk_text(r["transcript"])
        
        # Apply reconciled global speaker labels if available
        chunk_map = speaker_mappings.get(r["chunk_index"], {})
        reconciled_text = apply_speaker_mapping_to_chunk(cleaned_text, chunk_map)

        # Shift timestamps relative to whole audio start
        shifted_text = shift_timestamps_in_text(reconciled_text, r["start_sec"], chunk_index=r["chunk_index"])
        merged_sections.append(shifted_text)

    final_merged_transcript = "\n\n".join(merged_sections).strip()

    metadata = {
        "duration": format_timestamp(total_duration),
        "model": ", ".join(models_used_set) if models_used_set else model,
        "date": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    }

    docx_path = None
    if auto_save:
        docx_filename = f"{input_file.stem}.docx"
        docx_path = output_dir / docx_filename
        save_transcript_docx(
            output_path=docx_path,
            title=input_file.stem,
            merged_transcript=final_merged_transcript,
            metadata=metadata
        )
        log("\n" + "=" * 60)
        log(f"Merged docx saved successfully to: {docx_path}")
        log("=" * 60)

    log("\n" + "=" * 60)
    log(f"MERGED TRANSCRIPT PREVIEW ({len(merged_sections)} chunk(s) merged)")
    log("=" * 60 + "\n")
    log(final_merged_transcript)

    # Collect intermediate files generated for this run (master converted & chunk files)
    intermediate_files = []
    if master_audio.resolve() != input_file.resolve():
        intermediate_files.append(master_audio)
    for item in chunk_items:
        chunk_f = item[3]
        if chunk_f.resolve() != input_file.resolve() and chunk_f not in intermediate_files:
            intermediate_files.append(chunk_f)

    return {
        "transcript": final_merged_transcript,
        "total_duration": total_duration,
        "title": input_file.stem,
        "metadata": metadata,
        "suggested_filename": f"{input_file.stem}.docx",
        "output_dir": output_dir,
        "docx_path": docx_path,
        "intermediate_files": intermediate_files
    }

def main():
    parser = argparse.ArgumentParser(description="Split raw audio, transcribe concurrently, reconcile speakers, and output merged docx")
    parser.add_argument("audio_path", help="Path to input audio file")
    parser.add_argument("--model", type=str, default=DEFAULT_MODEL, help=f"Primary Gemini model to use (default: {DEFAULT_MODEL})")
    parser.add_argument("--chunk-minutes", type=str, default="auto", help="Target chunk length in minutes or 'auto' (default: auto)")
    parser.add_argument("--notes", action="store_true", help="Generate meeting notes and append to docx")
    parser.add_argument("--drive", action="store_true", help="Auto-upload to Google Drive")
    parser.add_argument("--max-workers", type=int, default=None, help="Maximum concurrent transcription workers (default: auto)")
    args = parser.parse_args()

    chunk_arg = args.chunk_minutes
    try:
        chunk_val = float(chunk_arg)
    except ValueError:
        chunk_val = "auto"

    try:
        res = run_transcription_pipeline(
            audio_path=args.audio_path,
            model=args.model,
            chunk_minutes=chunk_val,
            max_workers=args.max_workers,
            auto_save=False
        )

        meeting_notes = None
        topic_slug = ""
        company_name = ""
        person_names = []
        if args.notes:
            print("\nGenerating meeting notes from transcript...", flush=True)
            notes_res = generate_meeting_notes(res["transcript"], model=args.model)
            meeting_notes = notes_res.get("notes")
            topic_slug = notes_res.get("topic_slug", "")
            company_name = notes_res.get("company_name", "")
            meeting_type = notes_res.get("meeting_type", "")
            person_names = notes_res.get("person_names", [])
            print("\n=== MEETING NOTES ===\n", meeting_notes)
        else:
            meeting_type = ""

        rec_date = get_file_recording_date(args.audio_path)
        base_name = build_meeting_base_name(
            date_str=rec_date,
            company_name=company_name,
            meeting_type=meeting_type,
            person_names=person_names,
            topic_slug=topic_slug,
            fallback_title=res["title"]
        )
        suffix = "Full" if meeting_notes else "Transcript"
        suggested_name = f"{base_name}_{suffix}.docx"

        output_path = res["output_dir"] / suggested_name
        save_transcript_docx(
            output_path=output_path,
            title=res["title"],
            merged_transcript=res["transcript"],
            meeting_notes=meeting_notes,
            metadata=res["metadata"]
        )
        print(f"\n[SUCCESS] Document saved to: {output_path}")

        if args.drive:
            upload_to_google_drive(output_path)

    except Exception as e:
        print(f"\n[ERROR] {e}", file=sys.stderr)
        sys.exit(1)

if __name__ == "__main__":
    main()
